import os
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


def add_heading(document: Document, text: str, level: int = 1) -> None:
    document.add_heading(text, level=level)


def add_paragraph(document: Document, text: str) -> None:
    p = document.add_paragraph(text)
    p_format = p.paragraph_format
    p_format.space_after = Pt(6)


def add_bullet(document: Document, text: str) -> None:
    document.add_paragraph(text, style="List Bullet")


def add_numbered(document: Document, text: str) -> None:
    document.add_paragraph(text, style="List Number")


def add_table(document: Document, rows: int, cols: int, headers: list[str] | None = None) -> None:
    table = document.add_table(rows=rows, cols=cols)
    table.style = "Light List Accent 1"
    if headers:
        hdr_cells = table.rows[0].cells
        for idx, header in enumerate(headers):
            hdr_cells[idx].text = header


def build_document() -> Document:
    doc = Document()

    # Title
    title = doc.add_paragraph()
    run = title.add_run("Student Portal — Android + Django REST + MongoDB")
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
    run.font.size = Pt(20)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Version 1.0").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("\n")

    # 1. Overview
    add_heading(doc, "1. Overview", 1)
    add_paragraph(doc, (
        "This document provides comprehensive technical and product documentation for the Student "
        "Portal project. The system comprises an Android mobile application built with Kotlin and "
        "Jetpack Compose, and a backend powered by Django REST Framework with MongoDB for primary "
        "data storage. The app supports authentication, profile, academics, finance, library, "
        "virtual campus (Zoom/Google Meet), and a full voting system."
    ))

    # 2. Architecture
    add_heading(doc, "2. Architecture", 1)
    add_paragraph(doc, "High-level components:")
    add_bullet(doc, "Android App — Kotlin, Jetpack Compose, Navigation-Compose, Room, DataStore, Retrofit/OkHttp")
    add_bullet(doc, "Backend — Django 5, DRF, JWT (SimpleJWT), CORS, SQLite (Django) + MongoDB (domain data)")
    add_bullet(doc, "Data — MongoDB collections for users, profiles, academics, finance, voting; Room for offline cache")
    add_bullet(doc, "Video — Google Meet (Intent-based) integrated by default; Zoom SDK optional via gradle properties")

    # 3. Repositories & Modules
    add_heading(doc, "3. Repository Structure", 1)
    add_paragraph(doc, "Folders:")
    add_bullet(doc, "android/ — Android application module `app` with Compose UI and data layer")
    add_bullet(doc, "backend/ — Django project `student_api` with core app and MongoDB services")

    # 4. Environments
    add_heading(doc, "4. Environments & Dependencies", 1)
    add_paragraph(doc, "Android key versions:")
    add_bullet(doc, "AGP 8.9.1, Kotlin 2.0.21, Compose BOM 2024.09.00, compileSdk 35, minSdk 24")
    add_bullet(doc, "Room 2.6.1, DataStore 1.1.1, Retrofit 2.11.0, OkHttp 4.12.0, kotlinx-serialization 1.7.3")
    add_paragraph(doc, "Backend dependencies:")
    add_bullet(doc, "Django 5.1.1, DRF 3.15.2, CORS Headers 4.4.0, SimpleJWT 5.3.1, PyMongo 4.6.1, Argon2 23.1.0")

    # 5. Setup & Local Development
    add_heading(doc, "5. Setup & Local Development", 1)
    add_paragraph(doc, "Prerequisites: Android Studio Giraffe+, JDK 11, Python 3.10+, MongoDB (localhost:27017)")
    add_paragraph(doc, "Backend setup (Windows PowerShell):")
    add_bullet(doc, "python -m venv backend\\.venv")
    add_bullet(doc, "backend\\.venv\\Scripts\\Activate.ps1")
    add_bullet(doc, "pip install -r backend\\requirements.txt")
    add_bullet(doc, "cd backend && python manage.py migrate")
    add_bullet(doc, "python manage.py create_sample_data && python manage.py create_voting_indexes")
    add_bullet(doc, "python manage.py runserver 0.0.0.0:8000")
    add_paragraph(doc, "Android app:")
    add_bullet(doc, "Set base URL in `NetworkModule.kt` to one of: 10.0.2.2, <your-ip>, 127.0.0.1")
    add_bullet(doc, "Open project in Android Studio and Run on emulator/device")

    # 6. Configuration
    add_heading(doc, "6. Configuration", 1)
    add_paragraph(doc, "Android Manifest permissions: INTERNET, ACCESS_NETWORK_STATE, CAMERA, RECORD_AUDIO. Network security config allows cleartext for local dev.")
    add_paragraph(doc, "Backend settings: DEBUG=True, CORS allow all, JWT auth, SQLite for Django, MongoDB for domain collections.")

    # 7. Android App Details
    add_heading(doc, "7. Android Application", 1)
    add_paragraph(doc, "Key layers:")
    add_bullet(doc, "UI: Jetpack Compose screens for Login, Home, Profile, Academics, Finance, Library, Virtual Campus, Voting System")
    add_bullet(doc, "Data: `ApiService` (Retrofit), `NetworkModule` (OkHttp + JSON), repositories (`AuthRepository`, `ProfileRepository`, `VotingRepository`)")
    add_bullet(doc, "Storage: Room (Profile cache), DataStore (auth token)")
    add_paragraph(doc, "Networking:")
    add_bullet(doc, "Base URLs with fallback: http://10.0.2.2:8000/api/, http://<LAN-IP>:8000/api/, http://127.0.0.1:8000/api/")
    add_bullet(doc, "Auth token automatically attached via Authorization: Bearer <token>")
    add_paragraph(doc, "Video:")
    add_bullet(doc, "Google Meet integrated via Intent (opens in app if installed, otherwise browser)")
    add_bullet(doc, "Optional Zoom SDK via gradle properties ZOOM_SDK_KEY/SECRET/DOMAIN")

    # 8. Backend API
    add_heading(doc, "8. Backend API", 1)
    add_paragraph(doc, "Base Path: /api/")
    add_paragraph(doc, "Authentication:")
    add_bullet(doc, "POST /api/auth/jwt/login — obtain access/refresh tokens")
    add_bullet(doc, "POST /api/auth/jwt/refresh — refresh access token")
    add_paragraph(doc, "Custom Auth + Profile:")
    add_bullet(doc, "POST /api/auth/login — MongoDB-backed auth returning JWT via service")
    add_bullet(doc, "POST /api/auth/logout — no-op placeholder")
    add_bullet(doc, "GET /api/profile/me — Profile from MongoDB with default fallback")
    add_paragraph(doc, "Academics:")
    add_bullet(doc, "GET /api/academics/course-registration, /course-work, /exam-card, /exam-audit, /exam-result, /academic-leave, /clearance")
    add_paragraph(doc, "Finance:")
    add_bullet(doc, "GET /api/finance/fee-payment, /view-balance, /receipts (stubs)")
    add_paragraph(doc, "Virtual Campus:")
    add_bullet(doc, "GET /api/virtual/dashboard, /my-courses, /lectures")
    add_bullet(doc, "GET /api/virtual/zoom-rooms — list rooms")
    add_bullet(doc, "POST /api/virtual/zoom-rooms/{roomId}/attendance/mark — mark attendance")
    add_bullet(doc, "GET /api/virtual/zoom-rooms/{roomId}/attendance/summary — summary")
    add_paragraph(doc, "Voting System:")
    add_bullet(doc, "GET /api/voting/elections — list")
    add_bullet(doc, "POST /api/voting/elections — create { title }")
    add_bullet(doc, "GET /api/voting/elections/{id} — details")
    add_bullet(doc, "POST /api/voting/elections/{id}/candidates — add candidate { name }")
    add_bullet(doc, "POST /api/voting/elections/{id}/vote — cast vote { candidate }")
    add_bullet(doc, "GET /api/voting/elections/{id}/results — results list")

    # 9. Data Model (MongoDB)
    add_heading(doc, "9. Data Model (MongoDB)", 1)
    add_bullet(doc, "voting_elections: { _id, title, candidates: [ { name } ], results: { <candidate>: votes }, created_at, updated_at }")
    add_bullet(doc, "voting_votes: { election_id, user_id, candidate, created_at } with unique index (election_id, user_id)")
    add_bullet(doc, "attendance_records: { room_id, course_id, user_id, date, marked_at }")

    # 10. Security & Auth
    add_heading(doc, "10. Security & Auth", 1)
    add_bullet(doc, "JWT via SimpleJWT; Authorization: Bearer <token>")
    add_bullet(doc, "Password hashing with Argon2 (preferred)")
    add_bullet(doc, "CORS enabled for local development and emulator hosts")

    # 11. Observability
    add_heading(doc, "11. Observability & Logging", 1)
    add_bullet(doc, "Backend structured console logging with configurable LOG_LEVEL")
    add_bullet(doc, "Client network logs via OkHttp + Timber in DEBUG builds")

    # 12. Testing
    add_heading(doc, "12. Testing", 1)
    add_bullet(doc, "Android: unit tests under app/src/test and androidTest")
    add_bullet(doc, "Backend: Django tests and sample scripts (e.g., test_login.py)")

    # 13. Deployment
    add_heading(doc, "13. Deployment", 1)
    add_bullet(doc, "Backend: containerize with Python 3.10 base, configure env vars for MongoDB and SECRET_KEY")
    add_bullet(doc, "Android: configure base URL via build flavors or runtime settings for prod/staging")

    # 14. Roadmap
    add_heading(doc, "14. Roadmap", 1)
    add_bullet(doc, "Wire all Android screens to backend APIs (academics/finance/library)")
    add_bullet(doc, "Implement offline-first profile sync and error states in UI")
    add_bullet(doc, "Role-based access control and audit logging in backend")
    add_bullet(doc, "Pagination and filtering for list endpoints")
    add_bullet(doc, "Optional Zoom SDK native join; feature flag and fallback handling")
    add_bullet(doc, "CI/CD: GitHub Actions for lint, test, build, and release")

    # 15. Troubleshooting
    add_heading(doc, "15. Troubleshooting", 1)
    add_bullet(doc, "Emulator cannot reach backend: use http://10.0.2.2:8000/api/")
    add_bullet(doc, "401 on login: ensure MongoDB is running and sample data created")
    add_bullet(doc, "Vote conflict 409: unique constraint (election_id, user_id)")

    # 16. License
    add_heading(doc, "16. License", 1)
    add_paragraph(doc, "MIT (or update accordingly)")

    return doc


def main() -> None:
    project_root = Path(__file__).resolve().parents[2]
    output_dir = project_root / "docs"
    output_dir.mkdir(parents=True, exist_ok=True)
    output_path = output_dir / "StudentPortal_Documentation.docx"
    doc = build_document()
    doc.save(str(output_path))
    print(f"Documentation generated at: {output_path}")


if __name__ == "__main__":
    main()



